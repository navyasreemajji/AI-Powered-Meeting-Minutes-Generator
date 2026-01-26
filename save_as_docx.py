from docx import Document

def save_as_docx(minutes: dict, filename: str):
    doc = Document()

    # Title
    doc.add_heading("Meeting Minutes", level=0)

    # Summary
    doc.add_heading("Abstract Summary", level=1)
    doc.add_paragraph(minutes["abstract_summary"])

    # Transcript
    doc.add_heading("Transcript", level=1)
    doc.add_paragraph(minutes["transcript"])

    # Meeting Minutes (POINTS)
    doc.add_heading("Meeting Minutes", level=1)

    meeting_minutes = minutes["meeting_minutes"]

    # If meeting minutes are a list, write as numbered points
    if isinstance(meeting_minutes, list):
        for point in meeting_minutes:
            doc.add_paragraph(point, style="List Number")
    else:
        # fallback (safety)
        doc.add_paragraph(meeting_minutes)

    # Sentiment
    doc.add_heading("Sentiment", level=1)
    doc.add_paragraph(minutes["sentiment"])

    doc.save(filename)
